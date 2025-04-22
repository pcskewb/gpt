import openai, sys, platform, requests, streamlit as st
from openai import OpenAI, AssistantEventHandler
from PIL import Image
from io import BytesIO
from typing_extensions import override
import json
import base64
import pandas as pd
import docx
from dotenv import load_dotenv
import os

load_dotenv()

api_key = os.environ.get("api_key")
ID = os.environ.get("ID")

# ------------------------------------------------------------------
# 1) Setup OpenAI Client
# ------------------------------------------------------------------
bindings_version = getattr(openai, "__version__", "unknown")

ascii_user_agent = json.dumps(
    {
        "bindings_version": bindings_version,
        "lang": "python",
        "lang_version": sys.version.split()[0],
        "platform": platform.platform(),
        "publisher": "openai",
    },
    ensure_ascii=True
)

client = OpenAI(
    api_key=api_key,
    default_headers={"OpenAI-Client-User-Agent": ascii_user_agent}
)

assistant_id = ID

# ------------------------------------------------------------------
# 2) Streamlit UI Setup
# ------------------------------------------------------------------
with open("skewb-logomark 3.png", "rb") as image_file:
    encoded = base64.b64encode(image_file.read()).decode()

html_code = f"""
<h1 style='text-align: center; padding:0 px'>
    <img src='data:image/png;base64,{encoded}' alt='logo' style='height: 40px; vertical-align: middle; margin-right: 10px;'>
    SkewbGPT
</h1>
"""

st.set_page_config(page_title="🧠 SkewbAI Assistant")
st.markdown(html_code, unsafe_allow_html=True)

if "thread" not in st.session_state:
    st.session_state.thread = None

if "messages" not in st.session_state:
    st.session_state.messages = []

if "code_toggles" not in st.session_state:
    st.session_state.code_toggles = {}

# ------------------------------------------------------------------
# File Fetching
# ------------------------------------------------------------------
def fetch_file(file_id: str):
    url = f"https://api.openai.com/v1/files/{file_id}/content"
    resp = requests.get(url, headers={"Authorization": f"Bearer {client.api_key}"})
    if resp.ok:
        return BytesIO(resp.content)
    return None

# ------------------------------------------------------------------
# Toggle callback function
# ------------------------------------------------------------------
def toggle_code(key):
    st.session_state.code_toggles[key] = not st.session_state.code_toggles.get(key, False)

# ------------------------------------------------------------------
# Show Message History
# ------------------------------------------------------------------
for i, msg in enumerate(st.session_state.messages):
    with st.chat_message(msg["role"]):
        if "content" in msg:
            st.markdown(msg["content"])
        if "code" in msg and msg["code"]:
            # Generate a unique key for each message's toggle using both index and role
            unique_key = f"analyze_{msg['role']}_{i}"
            
            # Initialize toggle state if not exists
            if unique_key not in st.session_state.code_toggles:
                st.session_state.code_toggles[unique_key] = False
                
            # Display toggle and code
            analyze_toggle = st.toggle(
                "Analyze", 
                key=unique_key,
                value=st.session_state.code_toggles[unique_key],
                on_change=toggle_code,
                args=(unique_key,)
            )
            
            if st.session_state.code_toggles[unique_key]:
                st.code(msg["code"], language="python")
                
        if "image" in msg and msg["image"]:
            st.image(msg["image"])
        if "files" in msg and msg["files"]:
            for f in msg["files"]:
                if f["name"].endswith(".csv"):
                    st.dataframe(pd.read_csv(f["data"]))
                elif f["name"].endswith((".xls", ".xlsx")):
                    st.dataframe(pd.read_excel(f["data"]))
                elif f["name"].endswith((".doc", ".docx")):
                    doc = docx.Document(f["data"])
                    content = "\n".join([para.text for para in doc.paragraphs])
                    st.text_area("Word Document", content, height=200)
                else:
                    st.download_button(label=f"Download {f['name']}", data=f["data"], file_name=f["name"])

# ------------------------------------------------------------------
# Streamlit Chat Input
# ------------------------------------------------------------------
prompt = st.chat_input("Enter your message:")

# ------------------------------------------------------------------
# Custom Streamlit Handler
# ------------------------------------------------------------------

class StreamlitHandler(AssistantEventHandler):
    def __init__(self):
        super().__init__()
        self.text = ""
        self.code = ""
        self.images = []
        self.files = []
        self.code_ready = False

        self.txt_box = st.empty()
        self.code_container = st.container()

    @override
    def on_text_delta(self, delta, _):
        if delta.value is not None:
            self.text += delta.value
            self.txt_box.markdown(self.text + "▌")

    @override
    def on_tool_call_created(self, tc):
        if tc.type == "code_interpreter":
            snippet = tc.code_interpreter.input
            if not isinstance(snippet, str):
                snippet = snippet.value
            self.code += snippet

    @override
    def on_tool_call_delta(self, delta, _):
        if delta.type != "code_interpreter":
            return

        # Capture code during tool execution
        if delta.code_interpreter.input:
            snippet = delta.code_interpreter.input
            if not isinstance(snippet, str):
                snippet = snippet.value
            self.code += snippet

        # Show outputs if any
        for out in delta.code_interpreter.outputs or []:
            if out.type == "image":
                img_data = fetch_file(out.image.file_id)
                if img_data:
                    img = Image.open(img_data)
                    self.images.append(img)
                    st.image(img)
            elif out.type == "file":
                file_data = fetch_file(out.file.file_id)
                if file_data:
                    self.files.append({"name": out.file.filename, "data": file_data})
                    if out.file.filename.endswith(".csv"):
                        st.dataframe(pd.read_csv(file_data))
                    elif out.file.filename.endswith((".xls", ".xlsx")):
                        st.dataframe(pd.read_excel(file_data))
                    elif out.file.filename.endswith((".doc", ".docx")):
                        doc = docx.Document(file_data)
                        content = "\n".join([para.text for para in doc.paragraphs])
                        st.text_area("Word Document Content", content, height=200)
                    else:
                        st.download_button(label="Download File", data=file_data, file_name=out.file.filename)

# ------------------------------------------------------------------
# Process Input and Stream Response
# ------------------------------------------------------------------
if prompt:
    if st.session_state.thread is None:
        st.session_state.thread = client.beta.threads.create()

    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    client.beta.threads.messages.create(
        thread_id=st.session_state.thread.id,
        role="user",
        content=prompt
    )

    handler = StreamlitHandler()
    with st.chat_message("assistant"):
        with st.spinner("Thinking…"):
            with client.beta.threads.runs.stream(
                thread_id=st.session_state.thread.id,
                assistant_id=assistant_id,
                event_handler=handler
            ) as run_stream:
                run_stream.until_done()
        
        # Final text without cursor
        handler.txt_box.markdown(handler.text.strip())
        
        # Create a unique key for this response
        current_index = len(st.session_state.messages)
        unique_key = f"analyze_assistant_{current_index}"
        
        # Add toggle for code after run is complete
        if handler.code:
            # Initialize the toggle state in session state
            if unique_key not in st.session_state.code_toggles:
                st.session_state.code_toggles[unique_key] = False
                
            # Display toggle and code
            with handler.code_container:
                analyze_toggle = st.toggle(
                    "Analyze", 
                    key=unique_key,
                    value=st.session_state.code_toggles[unique_key],
                    on_change=toggle_code,
                    args=(unique_key,)
                )
                
                if st.session_state.code_toggles[unique_key]:
                    st.code(handler.code, language="python")

        # Save to message history
        st.session_state.messages.append({
            "role": "assistant",
            "content": handler.text.strip(),
            "code": handler.code if handler.code else None,
            "image": handler.images[0] if handler.images else None,
            "files": handler.files if handler.files else None
        })